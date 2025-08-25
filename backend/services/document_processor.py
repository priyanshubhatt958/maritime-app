import asyncio
import os
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime, timedelta
import docx
import PyPDF2
import aiofiles
from .ai_service import AIService
from .ocr_service import OCRService
from ..models.sof import SofEvent
from ..core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.ai_service = AIService()
        self.ocr_service = OCRService()
    
    async def process_sof_document(
        self, 
        file_path: str, 
        mode: str = "accuracy",
        port_timezone: str = "UTC",
        enable_ocr: bool = True
    ) -> Dict[str, Any]:
        """Process Statement of Facts document"""
        
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text = await self._extract_pdf_text(file_path, enable_ocr)
        elif file_ext in ['.docx', '.doc']:
            text = await self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        if not text.strip():
            raise ValueError("No text could be extracted from the document")
        
        # Process with AI for event extraction
        if mode == "accuracy":
            result = await self.ai_service.extract_sof_events(text, port_timezone)
        else:
            # Cost-saving mode - use simpler processing
            result = await self._extract_events_simple(text, port_timezone)
        
        # Post-process and validate events
        events = await self._validate_and_enhance_events(result.get('events', []))
        anomalies = result.get('anomalies', [])
        
        return {
            "events": events,
            "stats": {
                "total_events": len(events),
                "low_confidence_count": sum(1 for e in events if e.get('confidence', 0) < 0.85),
                "processing_time": datetime.utcnow().isoformat(),
                "text_length": len(text),
                "mode": mode
            },
            "anomalies": anomalies,
            "raw_text": text[:1000] + "..." if len(text) > 1000 else text
        }
    
    async def _extract_pdf_text(self, file_path: str, enable_ocr: bool = True) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            # Try native PDF text extraction first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            # If no text extracted and OCR is enabled, use OCR
            if not text.strip() and enable_ocr:
                text = await self.ocr_service.extract_text_from_pdf(file_path)
                
        except Exception as e:
            if enable_ocr:
                # Fallback to OCR
                text = await self.ocr_service.extract_text_from_pdf(file_path)
            else:
                raise Exception(f"PDF text extraction failed: {str(e)}")
        
        return text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    async def _extract_events_simple(self, text: str, port_timezone: str) -> Dict[str, Any]:
        """Simple event extraction without AI (cost-saving mode)"""
        import re
        
        events = []
        lines = text.split('\n')
        
        # Common maritime event patterns
        event_patterns = [
            r'arrived|arrival|berthed|anchored',
            r'commenced|started|began',
            r'completed|finished|ended',
            r'sailed|departed|left',
            r'loading|discharging|cargo',
            r'notice.*readiness|nor|n\.o\.r',
            r'pilot|tug|mooring',
            r'weather|delay|waiting'
        ]
        
        # Time patterns
        time_pattern = r'(\d{1,2}[:.]?\d{2})\s*(hrs?|hours?)?'
        date_pattern = r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            # Check if line contains event keywords
            for pattern in event_patterns:
                if re.search(pattern, line_lower):
                    # Extract time information
                    time_matches = re.findall(time_pattern, line)
                    date_matches = re.findall(date_pattern, line)
                    
                    if time_matches or date_matches:
                        events.append({
                            "event_name": line.strip()[:100],
                            "start_time_iso": datetime.utcnow().isoformat() + "Z",
                            "end_time_iso": None,
                            "duration_minutes": None,
                            "page": 1,
                            "row_index": i + 1,
                            "confidence": 0.7
                        })
                    break
        
        return {
            "events": events[:20],  # Limit to 20 events in simple mode
            "anomalies": []
        }
    
    async def _validate_and_enhance_events(self, events: List[Dict]) -> List[Dict]:
        """Validate and enhance extracted events"""
        enhanced_events = []
        
        for event in events:
            # Ensure required fields
            if not event.get('event_name'):
                continue
            
            # Validate and fix timestamps
            start_time = event.get('start_time_iso')
            end_time = event.get('end_time_iso')
            
            if start_time:
                try:
                    start_dt = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
                    event['start_time_iso'] = start_dt.isoformat() + 'Z'
                except:
                    event['start_time_iso'] = datetime.utcnow().isoformat() + 'Z'
            
            if end_time:
                try:
                    end_dt = datetime.fromisoformat(end_time.replace('Z', '+00:00'))
                    event['end_time_iso'] = end_dt.isoformat() + 'Z'
                    
                    # Calculate duration
                    if start_time:
                        start_dt = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
                        duration = (end_dt - start_dt).total_seconds() / 60
                        event['duration_minutes'] = int(duration) if duration > 0 else None
                except:
                    event['end_time_iso'] = None
            
            # Ensure confidence score
            if 'confidence' not in event:
                event['confidence'] = 0.8
            
            # Ensure page and row_index
            if 'page' not in event:
                event['page'] = 1
            if 'row_index' not in event:
                event['row_index'] = len(enhanced_events) + 1
            
            enhanced_events.append(event)
        
        return enhanced_events
    
    async def export_events_csv(self, events: List[Dict], file_path: str) -> str:
        """Export events to CSV format"""
        import csv
        
        with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = [
                'event_name', 'start_time_iso', 'end_time_iso', 
                'duration_minutes', 'page', 'row_index', 'confidence'
            ]
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            
            writer.writeheader()
            for event in events:
                writer.writerow(event)
        
        return file_path
    
    async def export_events_json(self, events: List[Dict], file_path: str) -> str:
        """Export events to JSON format"""
        export_data = {
            "events": events,
            "exported_at": datetime.utcnow().isoformat() + "Z",
            "total_events": len(events)
        }
        
        async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
            await f.write(json.dumps(export_data, indent=2, ensure_ascii=False))
        
        return file_path